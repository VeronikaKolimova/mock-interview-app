from pathlib import Path
import io
import zipfile
import logging

logger = logging.getLogger(__name__)

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Извлекает текст из содержимого файла с максимальной отказоустойчивостью.
    """
    extension = Path(filename).suffix.lower()
    logger.info(f"Попытка извлечь текст из файла: {filename} (расширение: {extension})")

    # 1. Обычный текст (самый надёжный)
    if extension == '.txt':
        try:
            return file_content.decode('utf-8', errors='ignore').strip()
        except Exception as e:
            raise ValueError(f"Ошибка чтения TXT: {str(e)}")

    # 2. PDF
    elif extension == '.pdf':
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if not text.strip():
                raise ValueError("PDF не содержит извлекаемого текста. Вероятно, это скан (изображение). Для таких файлов требуется OCR (распознавание текста), который пока не подключен. Пожалуйста, скопируйте текст вручную или сохраните файл как настоящий текстовый PDF.")
            
            return text.strip()
        except ImportError:
            raise ImportError("Для чтения PDF установите: pip install PyPDF2")
        except ValueError as ve:
            raise ValueError(str(ve))
        except Exception as e:
            logger.error(f"Ошибка чтения PDF: {str(e)}")
            raise ValueError(f"Ошибка чтения PDF. Возможно, файл повреждён или защищён паролем. Детали: {str(e)}")

    # 3. DOCX (и попытка поймать поддельные .docx)
    elif extension in ['.docx', '.doc']:
        try:
            import docx
            # Проверка на ZIP-архив (настоящий .docx это ZIP)
            try:
                with zipfile.ZipFile(io.BytesIO(file_content)) as z:
                    z.testzip()
            except zipfile.BadZipFile:
                raise ValueError("Файл имеет расширение .docx/.doc, но не является корректным архивом Word. Возможно, это старый формат .doc или файл сохранён из веб-страницы. Откройте его в Word и выберите 'Сохранить как' -> 'Документ Word (*.docx)' или 'PDF'.")
            
            doc = docx.Document(io.BytesIO(file_content))
            
            # Собираем текст из параграфов
            text_parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            # ДОПОЛНИТЕЛЬНО: пытаемся вытащить текст из таблиц (частая проблема docx)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            final_text = "\n".join(text_parts)
            
            if not final_text:
                raise ValueError("Не удалось найти текст в документе. Возможно, текст находится в текстовых полях (Text Boxes) или колонтитулах, которые python-docx не извлекает. Попробуйте скопировать текст вручную или сохранить как PDF.")
                
            return final_text
            
        except ImportError:
            raise ImportError("Для чтения DOCX установите: pip install python-docx")
        except ValueError as ve:
            raise ValueError(str(ve))
        except Exception as e:
            logger.error(f"Критическая ошибка чтения DOCX: {str(e)}")
            raise ValueError(f"Не удалось прочитать документ из-за внутренней ошибки: {str(e)}. Попробуйте пересохранить файл.")

    else:
        raise ValueError(f"Формат '{extension}' не поддерживается. Пожалуйста, используйте .txt, .pdf или .docx")
